# 导入所需的库
import os
from PyPDF2 import PdfReader
import docx
from typing import List, IO, Any

# 导入 LangChain 相关的核心类
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_deepseek import ChatDeepSeek
from langchain.chains import ConversationalRetrievalChain

# (可选) 如果你还没有设置环境变量，可以在这里加载
from dotenv import load_dotenv
load_dotenv()


def get_documents_from_files(uploaded_files: List[IO[Any]]) -> List[Document]:
    """
    从上传的文件对象列表中提取文本并创建 LangChain Document 对象。

    Args:
        uploaded_files: 一个文件对象列表。每个对象需要有 .name 和 .read() 方法。

    Returns:
        一个 Document 对象的列表。
    """
    docs = []
    for file_io in uploaded_files:
        # 为了模拟 Streamlit 的 UploadedFile 对象，我们假设它有 .name 属性
        file_name = getattr(file_io, 'name', 'unknown_file')

        if file_name.lower().endswith(".pdf"):
            try:
                pdf_reader = PdfReader(file_io)
                for idx, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        docs.append(
                            Document(
                                page_content=page_text,
                                metadata={"source": f"{file_name} on page {idx + 1}"}
                            )
                        )
            except Exception as e:
                print(f"Error reading PDF file {file_name}: {e}")

        elif file_name.lower().endswith(".docx"):
            try:
                doc = docx.Document(file_io)
                for idx, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text:
                        docs.append(
                            Document(
                                page_content=paragraph.text,
                                metadata={"source": f"{file_name} in paragraph {idx + 1}"}
                            )
                        )
            except Exception as e:
                print(f"Error reading DOCX file {file_name}: {e}")

        elif file_name.lower().endswith(".txt"):
            try:
                text = file_io.read().decode("utf-8")
                if text:
                    docs.append(Document(page_content=text, metadata={"source": file_name}))
            except Exception as e:
                print(f"Error reading TXT file {file_name}: {e}")

    # 注意：这里的 return 语句在 for 循环之外，是修正后的正确位置。
    return docs


def get_text_chunks(docs: List[Document]) -> List[Document]:
    """将文档分割成小块文本"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    docs_chunks = text_splitter.split_documents(docs)
    return docs_chunks


def get_vectorstore(text_chunks: List[Document]):
    """创建向量存储 (使用阿里云灵积)"""
    if not os.getenv("DASHSCOPE_API_KEY"):
        raise ValueError("DASHSCOPE_API_KEY environment variable not set.")

    embeddings = DashScopeEmbeddings(
        model="text-embedding-v3",
        dashscope_api_key=os.getenv("DASHSCOPE_API_KEY")
    )
    vectorstore = FAISS.from_documents(text_chunks, embedding=embeddings)
    return vectorstore


def get_conversation_chain(vectorstore):
    """创建对话检索链 (使用深度求索)"""
    if not os.getenv("DEEPSEEK_API_KEY"):
        raise ValueError("DEEPSEEK_API_KEY environment variable not set.")

    llm = ChatDeepSeek(
        model="deepseek-chat",  # 使用正确的 DeepSeek 模型名称
        api_key=os.getenv("DEEPSEEK_API_KEY"),
        temperature=0,
    )
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        return_source_documents=True,
    )
    return conversation_chain


def ask_question(chain, question: str, chat_history: list) -> dict:
    """
    使用对话链进行提问，并返回结果。
    """
    response = chain.invoke({"question": question, "chat_history": chat_history})
    return response


# --- 主执行逻辑 ---
if __name__ == "__main__":
    # 这是一个示例，演示如何独立运行这个后端逻辑

    # 1. 定义你的文件路径
    file_paths = [
        "D:/AIProject/jvke/day11-demo/pdf-search/resource/chatgpt.pdf",
        "D:/AIProject/jvke/day11-demo/pdf-search/resource/knowledge.txt",
        "D:/AIProject/jvke/day11-demo/pdf-search/resource/labor-law.pdf"
    ]

    print("--- 1. Reading and parsing documents ---")
    # 2. 模拟文件上传：打开文件并创建文件IO对象列表
    try:
        uploaded_files_simulation = [open(path, "rb") for path in file_paths]

        # 3. 提取文档内容
        documents = get_documents_from_files(uploaded_files_simulation)
        if not documents:
            print("No documents were parsed. Please check file paths and content.")
        else:
            print(f"Successfully parsed {len(documents)} document sections.")

            # 4. 分割文本
            text_chunks = get_text_chunks(documents)
            print(f"Split documents into {len(text_chunks)} chunks.")

            # 5. 创建向量存储
            print("\n--- 2. Creating vector store using DashScope (this may take a while) ---")
            vector_store = get_vectorstore(text_chunks)
            print("Vector store created successfully.")

            # 6. 创建对话链
            conversation_chain = get_conversation_chain(vector_store)
            print("Conversation chain with DeepSeek created successfully.")

            # 7. 开始交互式问答循环
            print("\n--- 3. Ready to chat! (Type 'exit' to quit) ---")
            chat_history = []
            while True:
                user_question = input("You: ")
                if user_question.lower() == 'exit':
                    break

                # 提问
                result = ask_question(conversation_chain, user_question, chat_history)
                answer = result["answer"]

                # 更新对话历史
                chat_history.append(("user", user_question))
                chat_history.append(("assistant", answer))

                print(f"Bot: {answer}")

                # 打印来源
                sources = result.get("source_documents", [])
                if sources:
                    source_names = set([doc.metadata.get("source", "Unknown") for doc in sources])
                    print(f"> Sources: {', '.join(source_names)}")

                print("-" * 20)

    except FileNotFoundError as e:
        print(f"Error: Could not find a file. Please check your file paths. Details: {e}")
    except ValueError as e:
        print(f"Configuration Error: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    finally:
        # 关闭所有打开的文件
        if 'uploaded_files_simulation' in locals():
            for f in uploaded_files_simulation:
                f.close()